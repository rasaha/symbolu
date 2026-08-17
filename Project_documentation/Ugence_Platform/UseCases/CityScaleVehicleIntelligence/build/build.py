import os, html, cairosvg
from content import CONTENT
import diagrams as DG

BASE='UGENCE_CITY_SCALE_VEHICLE_INTELLIGENCE_CLIENT_BRIEF_v2.0'
META=dict(
 title='Ugence Governance for City-Scale Vehicle Intelligence — Client Architecture, Governance and Execution-Assurance Briefing',
 author='Ugence Labs',
 company='Ugence Labs',
 subject='Conceptual reference architecture for governing city-scale vehicle intelligence with the current Ugence Labs portfolio',
 keywords='Ugence, vehicle intelligence, ANPR, ALPR, vehicle re-identification, trajectory, AI governance, ActionGate, Action Clearance, Risk Authority, Decision Authority, Agent Runtime, execution assurance, effect verification, smart city, public sector')

os.makedirs('svg',exist_ok=True); os.makedirs('png',exist_ok=True)
SVGS={}
for k,fn in DG.ALL.items():
    svg,w,h=fn(); SVGS[k]=(svg,w,h)
    open(f'svg/{k}.svg','w').write(svg)
    cairosvg.svg2png(bytestring=svg.encode(), write_to=f'png/{k}.png', scale=2.2, background_color='white')

def esc(t): return html.escape(str(t))
def rhtml(r):
    if isinstance(r,str): return esc(r)
    t=esc(r['text'])
    if r.get('b'): t=f'<strong>{t}</strong>'
    if r.get('i'): t=f'<em>{t}</em>'
    return t
def runs_html(runs): return ''.join(rhtml(r) for r in runs)
def item_html(it): return runs_html(it) if isinstance(it,list) else esc(it)

# --------- collect TOC entries (h1 + h2 with id) ----------
TOC=[]
for blk in CONTENT:
    if blk['t']=='h1': TOC.append((1,blk['id'],blk['text']))
    elif blk['t']=='h2' and blk.get('id'): TOC.append((2,blk['id'],blk['text']))

# ======================= HTML / PDF =======================
def build_html():
    css='''
@page { size:A4 portrait; margin:20mm 17mm 16mm 17mm;
 @top-left{ content:"Ugence · City-Scale Vehicle Intelligence"; font:8pt Arial; color:#9aa0b0;}
 @top-right{ content:"Client Briefing · v2.0"; font:8pt Arial; color:#9aa0b0;}
 @bottom-left{ content:"Ugence Labs · ugence.ai"; font:8pt Arial; color:#9aa0b0;}
 @bottom-right{ content:"Page " counter(page) " of " counter(pages); font:8pt Arial; color:#9aa0b0;} }
@page cover { margin:0; @top-left{content:none} @top-right{content:none} @bottom-left{content:none} @bottom-right{content:none} }
@page toc:first { @top-left{content:none} @top-right{content:none} }
html{ font-family:Arial,Helvetica,sans-serif; font-size:10.5pt; color:#1f2430; line-height:1.42; }
h1{ font-family:Georgia,'Times New Roman',serif; font-size:19pt; color:#1a1740; margin:16pt 0 8pt; break-after:avoid; }
h1:first-of-type{ margin-top:0; }
ul.toc{ list-style:none; padding-left:0; }
ul.toc li{ margin:4pt 0; }
a.tl{ text-decoration:none; color:#1f2430; }
a.tl::after{ content: leader('.') " " target-counter(attr(href), page); color:#667; font-weight:normal; }
h2{ font-family:Georgia,serif; font-size:14pt; color:#2A2170; margin:14pt 0 5pt; break-after:avoid; }
h3{ font-size:11.5pt; color:#2A2170; margin:10pt 0 4pt; break-after:avoid; }
p{ margin:0 0 7pt; }
ul,ol{ margin:0 0 7pt; padding-left:18pt; }
li{ margin:0 0 3pt; }
strong{ color:#111; }
table{ width:100%; border-collapse:collapse; margin:4pt 0 2pt; font-size:9.3pt; break-inside:auto; }
thead{ display:table-header-group; }
th{ background:#1E2340; color:#fff; text-align:left; padding:5pt 6pt; font-size:9.3pt; }
td{ border:0.5pt solid #d8dee9; padding:4pt 6pt; vertical-align:top; }
tbody tr:nth-child(even){ background:#f5f7fb; }
.caption{ font-size:8.6pt; font-style:italic; color:#667; margin:2pt 0 10pt; }
figure{ margin:6pt 0 2pt; text-align:center; break-inside:avoid; }
figure img{ max-width:100%; }
.callout{ border:0.75pt solid #cfd6e6; border-left:3pt solid #5145C7; background:#f4f5fb; padding:7pt 10pt; margin:8pt 0; border-radius:3pt; break-inside:avoid; font-size:9.8pt; }
.callout.warn{ border-left-color:#B7791F; background:#fbf6ec; }
.callout.note{ border-left-color:#2E7D5B; background:#eef6f1; }
.callout .ct{ font-weight:bold; color:#2A2170; display:block; margin-bottom:2pt; }
.callout.warn .ct{ color:#8a5a12; } .callout.note .ct{ color:#1c5c40; }
.quote{ font-family:Georgia,serif; font-size:12pt; font-style:italic; color:#2A2170; border-left:3pt solid #5145C7; padding:2pt 0 2pt 12pt; margin:10pt 0; break-inside:avoid; }
.modcard{ border:0.75pt solid #d8dee9; border-radius:4pt; margin:7pt 0; break-inside:avoid; }
.modcard .mh{ background:#ECEAFB; color:#2A2170; font-weight:bold; padding:5pt 8pt; font-size:10.5pt; display:flex; justify-content:space-between; }
.modcard .badge{ font-size:7.6pt; font-weight:bold; color:#fff; background:#5145C7; border-radius:8pt; padding:2pt 7pt; }
.modcard table{ margin:0; } .modcard td{ border:none; border-top:0.5pt solid #eef; }
.modcard td.k{ width:26%; font-weight:bold; color:#374151; background:#fafbfe; }
.uc{ border:0.75pt solid #d8dee9; border-radius:4pt; margin:6pt 0; padding:0; break-inside:avoid; }
.uc .uh{ background:#f0f2fa; color:#2A2170; font-weight:bold; padding:5pt 8pt; font-size:10.6pt; }
.uc table{ margin:0; font-size:9pt; } .uc td{ border:none; border-top:0.5pt solid #eef; }
.uc td.k{ width:22%; font-weight:bold; color:#374151; }
.uc .badge{ font-size:7.6pt; font-weight:bold; color:#fff; background:#5145C7; border-radius:8pt; padding:2pt 7pt; }
.uc .badge.future{ background:#D97706; }
'''
    parts=['<!doctype html><html><head><meta charset="utf-8">',
      f'<title>{esc(META["title"])}</title>',
      f'<meta name="author" content="{esc(META["author"])}">',
      f'<meta name="description" content="{esc(META["subject"])}">',
      f'<meta name="keywords" content="{esc(META["keywords"])}">',
      f'<style>{css}</style></head><body>']
    # cover
    parts.append('''<section style="page:cover; height:297mm; box-sizing:border-box; background:#16143A; color:#fff; padding:34mm 22mm;">
    <div style="font:11pt Arial; letter-spacing:2px; color:#9aa0d8; text-transform:uppercase;">Ugence Labs Portfolio · Smart City &amp; Public-Sector Solutions</div>
    <div style="height:16mm"></div>
    <div style="font-family:Georgia,serif; font-size:34pt; font-weight:bold; line-height:1.08;">Ugence Governance for<br>City-Scale Vehicle Intelligence</div>
    <div style="height:6mm"></div>
    <div style="font-family:Georgia,serif; font-size:18pt; color:#ECEAFB;">Client Architecture, Governance and Execution-Assurance Briefing</div>
    <div style="height:14mm"></div>
    <div style="font-size:12.5pt; color:#c9cfea; max-width:152mm; line-height:1.5;">The vehicle-intelligence platform determines what the evidence indicates. Ugence determines what may be concluded, disclosed, investigated or executed — under which policy, evidence, authority, purpose and scope — and verifies that the resulting action remained within that authorization.</div>
    <div style="height:24mm"></div>
    <div style="font-size:13pt;"><strong style="color:#fff;">Prepared by Ugence Labs</strong></div>
    <div style="font-size:11.5pt; color:#c9cfea;">ugence.ai · Version 2.0 · August 2026</div>
    <div style="height:18mm"></div>
    <div style="font-size:9.5pt; font-style:italic; color:#aeb6e0; max-width:158mm; line-height:1.5;">Conceptual reference architecture grounded in the current Ugence repository. Not legal advice, not a representation of an existing deployment, and not a substitute for jurisdiction-specific policy review. Capability maturity is stated conservatively throughout.</div>
    </section>''')
    # TOC
    parts.append('<section style="page:toc; break-before:page; break-after:page;"><h1 style="break-before:avoid; margin-top:0;">Contents</h1><ul class="toc">')
    for lvl,idv,txt in TOC:
        pad='0' if lvl==1 else '16pt'
        wt='bold' if lvl==1 else 'normal'
        parts.append(f'<li style="margin-left:{pad}; font-weight:{wt};"><a class="tl" href="#{idv}">{esc(txt)}</a></li>')
    parts.append('</ul></section>')
    # body
    for blk in CONTENT:
        t=blk['t']
        if t=='h1': parts.append(f'<h1 id="{blk["id"]}">{esc(blk["text"])}</h1>')
        elif t=='h2':
            idattr=f' id="{blk["id"]}"' if blk.get('id') else ''
            parts.append(f'<h2{idattr}>{esc(blk["text"])}</h2>')
        elif t=='h3': parts.append(f'<h3>{esc(blk["text"])}</h3>')
        elif t=='p': parts.append(f'<p>{runs_html(blk["runs"])}</p>')
        elif t=='bul': parts.append('<ul>'+''.join(f'<li>{item_html(it)}</li>' for it in blk['items'])+'</ul>')
        elif t=='num': parts.append('<ol>'+''.join(f'<li>{item_html(it)}</li>' for it in blk['items'])+'</ol>')
        elif t=='quote': parts.append(f'<div class="quote">{esc(blk["text"])}</div>')
        elif t=='callout':
            ct=f'<span class="ct">{esc(blk["title"])}</span>' if blk.get('title') else ''
            parts.append(f'<div class="callout {blk["kind"]}">{ct}{esc(blk["text"])}</div>')
        elif t=='pagebreak': parts.append('<div style="break-after:page"></div>')
        elif t=='diagram':
            svg=SVGS[blk['svg']][0]
            parts.append(f'<figure>{svg}<figcaption class="caption">{esc(blk["caption"])}</figcaption></figure>')
        elif t=='table':
            ws=blk.get('widths')
            colg=''
            if ws: colg='<colgroup>'+''.join(f'<col style="width:{w*100:.1f}%">' for w in ws)+'</colgroup>'
            th='<thead><tr>'+''.join(f'<th>{esc(h)}</th>' for h in blk['headers'])+'</tr></thead>'
            body='<tbody>'+''.join('<tr>'+''.join(f'<td>{esc(c)}</td>' for c in row)+'</tr>' for row in blk['rows'])+'</tbody>'
            parts.append(f'<table>{colg}{th}{body}</table><div class="caption">{esc(blk["caption"])}</div>')
        elif t=='modcard':
            rows=''.join(f'<tr><td class="k">{esc(k)}</td><td>{esc(v)}</td></tr>' for k,v in blk['rows'])
            parts.append(f'<div class="modcard"><div class="mh"><span>{esc(blk["name"])}</span><span class="badge">{esc(blk["maturity"])}</span></div><table>{rows}</table></div>')
        elif t=='uccards':
            for c in blk['cards']:
                fut=' future' if 'FUTURE' in c['maturity'] else ''
                rows=''.join(f'<tr><td class="k">{esc(k)}</td><td>{esc(v)}</td></tr>' for k,v in
                    [('Business problem',c['problem']),('ServiceNow anchor',c['anchor']),('Ugence proposes',c['ugence']),('Autonomy boundary',c['autonomy']),('Discovery question',c['discovery'])])
                parts.append(f'<div class="uc"><div class="uh">{esc(c["id"])} · {esc(c["name"])} &nbsp; <span class="badge{fut}">{esc(c["maturity"])}</span></div><table>{rows}</table></div>')
    parts.append('</body></html>')
    return ''.join(parts)

def render_pdf():
    from weasyprint import HTML
    HTML(string=build_html()).write_pdf(BASE+'.pdf')
    # clean metadata
    import pikepdf
    pdf=pikepdf.open(BASE+'.pdf', allow_overwriting_input=True)
    di=pdf.docinfo
    di[pikepdf.Name.Title]=META['title']; di[pikepdf.Name.Author]=META['author']
    di[pikepdf.Name.Subject]=META['subject']; di[pikepdf.Name.Keywords]=META['keywords']
    di[pikepdf.Name.Creator]='Ugence Labs'; di[pikepdf.Name.Producer]='Ugence Labs'
    with pdf.open_metadata() as m:
        m['dc:title']=META['title']; m['dc:creator']=[META['author']]
        m['dc:description']=META['subject']; m['pdf:Keywords']=META['keywords']; m['xmp:CreatorTool']='Ugence Labs'
    pdf.save(BASE+'.pdf'); pdf.close()
    print('PDF written')

# ======================= DOCX =======================
def render_docx():
    from docx import Document
    from docx.shared import Pt, Mm, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    doc=Document()
    # A4 portrait + margins
    sec=doc.sections[0]
    sec.page_width=Mm(210); sec.page_height=Mm(297)
    sec.left_margin=Mm(20); sec.right_margin=Mm(18); sec.top_margin=Mm(18); sec.bottom_margin=Mm(16)
    sec.header_distance=Mm(10); sec.footer_distance=Mm(9)
    # base style
    st=doc.styles['Normal']; st.font.name='Arial'; st.font.size=Pt(10.5)
    def set_field(run, instr):
        f1=OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'),'begin')
        it=OxmlElement('w:instrText'); it.set(qn('xml:space'),'preserve'); it.text=instr
        f2=OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'),'end')
        run._r.append(f1); run._r.append(it); run._r.append(f2)
    def shade(cell,color):
        tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:fill'),color); tcPr.append(shd)
    def set_cell_w(cell,w_mm):
        cell.width=Mm(w_mm)
    def no_border_cover(): pass
    # header/footer (skip on first/cover via different first page)
    sec.different_first_page_header_footer=True
    hp=sec.header.paragraphs[0]; hp.text=''; r=hp.add_run('Ugence Governance for City-Scale Vehicle Intelligence — Client Briefing')
    r.font.size=Pt(8); r.font.color.rgb=RGBColor(0x9a,0xa0,0xb0)
    fp=sec.footer.paragraphs[0]; fp.text=''
    fr=fp.add_run('Ugence Labs · ugence.ai · Client Briefing v2.0      Page '); fr.font.size=Pt(8); fr.font.color.rgb=RGBColor(0x9a,0xa0,0xb0)
    pr=fp.add_run(); pr.font.size=Pt(8); pr.font.color.rgb=RGBColor(0x9a,0xa0,0xb0); set_field(pr,'PAGE')
    # ---- cover ----
    def cover_par(text,size,color,bold=False,italic=False,after=4):
        p=doc.add_paragraph(); run=p.add_run(text); run.font.size=Pt(size); run.bold=bold; run.italic=italic
        run.font.color.rgb=color; p.paragraph_format.space_after=Pt(after); return p
    NAVY=RGBColor(0x1a,0x17,0x40); VIOL=RGBColor(0x2A,0x21,0x70); GREY=RGBColor(0x55,0x5b,0x66)
    for _ in range(2): doc.add_paragraph()
    cover_par('UGENCE LABS PORTFOLIO · SMART CITY & PUBLIC-SECTOR SOLUTIONS',10,RGBColor(0x51,0x45,0xC7),bold=True,after=10)
    cover_par('Ugence Governance for City-Scale Vehicle Intelligence',30,NAVY,bold=True,after=6)
    cover_par('Client Architecture, Governance and Execution-Assurance Briefing',16,VIOL,bold=True,after=12)
    cover_par('The vehicle-intelligence platform determines what the evidence indicates. Ugence determines what may be concluded, disclosed, investigated or executed — under which policy, evidence, authority, purpose and scope — and verifies that the resulting action remained within that authorization.',12,RGBColor(0x37,0x41,0x51),after=18)
    cover_par('Prepared by Ugence Labs',13,NAVY,bold=True,after=1)
    cover_par('ugence.ai · Version 2.0 · August 2026',11,RGBColor(0x37,0x41,0x51),after=18)
    cover_par('Conceptual reference architecture grounded in the current Ugence repository. Not legal advice, not a representation of an existing deployment, and not a substitute for jurisdiction-specific policy review. Capability maturity is stated conservatively throughout.',9.5,GREY,italic=True,after=2)
    doc.add_page_break()
    # ---- TOC ----
    h=doc.add_heading('Contents',level=1)
    ptoc=doc.add_paragraph(); run=ptoc.add_run()
    f1=OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'),'begin')
    it=OxmlElement('w:instrText'); it.set(qn('xml:space'),'preserve'); it.text='TOC \\o "1-2" \\h \\z \\u'
    f2=OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'),'separate')
    tt=OxmlElement('w:t'); tt.text='Right-click and choose “Update Field” to build the table of contents.'
    f3=OxmlElement('w:fldChar'); f3.set(qn('w:fldCharType'),'end')
    run._r.append(f1); run._r.append(it); run._r.append(f2); run._r.append(tt); run._r.append(f3)
    doc.add_page_break()
    # update fields on open
    settings=doc.settings.element
    uf=OxmlElement('w:updateFields'); uf.set(qn('w:val'),'true'); settings.append(uf)
    # ---- body ----
    def add_runs(p, runs):
        for r in runs:
            if isinstance(r,str): p.add_run(r)
            else:
                run=p.add_run(r['text']);
                if r.get('b'): run.bold=True
                if r.get('i'): run.italic=True
    def caption(text):
        p=doc.add_paragraph(); r=p.add_run(text); r.italic=True; r.font.size=Pt(8.5); r.font.color.rgb=RGBColor(0x66,0x66,0x77)
        p.paragraph_format.space_after=Pt(8)
    CONTENT_W=210-20-18  # mm usable ~172
    for blk in CONTENT:
        t=blk['t']
        if t=='h1': doc.add_heading(blk['text'],level=1)
        elif t=='h2': doc.add_heading(blk['text'],level=2)
        elif t=='h3': doc.add_heading(blk['text'],level=3)
        elif t=='p':
            p=doc.add_paragraph(); add_runs(p,blk['runs'])
        elif t=='bul':
            for it2 in blk['items']:
                p=doc.add_paragraph(style='List Bullet')
                if isinstance(it2,list): add_runs(p,it2)
                else: p.add_run(it2)
        elif t=='num':
            for it2 in blk['items']:
                p=doc.add_paragraph(style='List Number')
                if isinstance(it2,list): add_runs(p,it2)
                else: p.add_run(it2)
        elif t=='quote':
            tbl=doc.add_table(rows=1,cols=1); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
            c=tbl.cell(0,0); shade(c,'F1F0FB')
            pp=c.paragraphs[0]; r=pp.add_run(blk['text']); r.italic=True; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0x2A,0x21,0x70)
            doc.add_paragraph().paragraph_format.space_after=Pt(2)
        elif t=='callout':
            fill={'note':'EEF6F1','warn':'FBF6EC','note2':'F4F5FB'}.get(blk['kind'],'F4F5FB')
            tbl=doc.add_table(rows=1,cols=1); c=tbl.cell(0,0); shade(c,fill)
            pp=c.paragraphs[0]
            if blk.get('title'):
                rr=pp.add_run(blk['title']+'  '); rr.bold=True; rr.font.color.rgb=RGBColor(0x2A,0x21,0x70)
            pp.add_run(blk['text']).font.size=Pt(9.8)
            doc.add_paragraph().paragraph_format.space_after=Pt(2)
        elif t=='pagebreak': doc.add_page_break()
        elif t=='diagram':
            png=f'png/{blk["svg"]}.png'
            p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(png, width=Mm(168))
            caption(blk['caption'])
        elif t=='table':
            headers=blk['headers']; rows=blk['rows']; ws=blk.get('widths')
            tbl=doc.add_table(rows=1,cols=len(headers)); tbl.style='Table Grid'; tbl.autofit=False
            tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
            hdr=tbl.rows[0].cells
            for i,htext in enumerate(headers):
                shade(hdr[i],'1E2340'); pp=hdr[i].paragraphs[0]; r=pp.add_run(htext); r.bold=True; r.font.color.rgb=RGBColor(0xff,0xff,0xff); r.font.size=Pt(9.2)
            for ri,row in enumerate(rows):
                cells=tbl.add_row().cells
                for ci,val in enumerate(row):
                    if ri%2==1: shade(cells[ci],'F5F7FB')
                    pp=cells[ci].paragraphs[0]; r=pp.add_run(str(val)); r.font.size=Pt(9.0)
            if ws:
                for ci,w in enumerate(ws):
                    wmm=CONTENT_W*w
                    for row in tbl.rows: row.cells[ci].width=Mm(wmm)
            caption(blk['caption'])
        elif t=='modcard':
            hp=doc.add_paragraph(); r=hp.add_run(blk['name']); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=RGBColor(0x2A,0x21,0x70)
            r2=hp.add_run('    ['+blk['maturity']+']'); r2.font.size=Pt(8); r2.bold=True; r2.font.color.rgb=RGBColor(0x51,0x45,0xC7)
            tbl=doc.add_table(rows=0,cols=2); tbl.style='Table Grid'
            for k,v in blk['rows']:
                cells=tbl.add_row().cells; shade(cells[0],'FAFBFE')
                rr=cells[0].paragraphs[0].add_run(k); rr.bold=True; rr.font.size=Pt(9)
                cells[1].paragraphs[0].add_run(v).font.size=Pt(9)
                cells[0].width=Mm(CONTENT_W*0.28); cells[1].width=Mm(CONTENT_W*0.72)
            doc.add_paragraph().paragraph_format.space_after=Pt(2)
        elif t=='uccards':
            for c in blk['cards']:
                hp=doc.add_paragraph(); r=hp.add_run(f'{c["id"]} · {c["name"]}'); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=RGBColor(0x2A,0x21,0x70)
                badge=hp.add_run('    ['+c['maturity']+']'); badge.font.size=Pt(8); badge.bold=True
                badge.font.color.rgb=RGBColor(0xD9,0x77,0x06) if 'FUTURE' in c['maturity'] else RGBColor(0x51,0x45,0xC7)
                tbl=doc.add_table(rows=0,cols=2); tbl.style='Table Grid'
                for k,v in [('Business problem',c['problem']),('ServiceNow anchor',c['anchor']),('Ugence proposes',c['ugence']),('Autonomy boundary',c['autonomy']),('Discovery question',c['discovery'])]:
                    cells=tbl.add_row().cells; shade(cells[0],'FAFBFE')
                    rr=cells[0].paragraphs[0].add_run(k); rr.bold=True; rr.font.size=Pt(9)
                    cells[1].paragraphs[0].add_run(v).font.size=Pt(9)
                    cells[0].width=Mm(CONTENT_W*0.24); cells[1].width=Mm(CONTENT_W*0.76)
                doc.add_paragraph().paragraph_format.space_after=Pt(2)
    # metadata
    cp=doc.core_properties
    cp.title=META['title']; cp.author=META['author']; cp.subject=META['subject']
    cp.keywords=META['keywords']; cp.category='Architecture briefing'; cp.comments=''; cp.last_modified_by='Rakesh Mohan'
    doc.save(BASE+'.docx')
    # set Company (extended property) in docProps/app.xml
    import zipfile
    src=BASE+'.docx'
    app=zipfile.ZipFile(src).read('docProps/app.xml').decode('utf8')
    if '<Company>' not in app:
        app=app.replace('</Properties>','<Company>Ugence Labs</Company></Properties>')
    else:
        import re as _re; app=_re.sub(r'<Company>.*?</Company>','<Company>Ugence Labs</Company>',app)
    tmp=src+'.tmp'; zin=zipfile.ZipFile(src); zout=zipfile.ZipFile(tmp,'w',zipfile.ZIP_DEFLATED)
    for it in zin.namelist():
        zout.writestr(it, app if it=='docProps/app.xml' else zin.read(it))
    zin.close(); zout.close(); os.replace(tmp,src)
    print('DOCX written')

if __name__=='__main__':
    render_pdf()
    render_docx()
    print('done')
