"""Build the Ugence Labs .docx (and .pdf) edition of the Capability Pipeline document.

Source of truth: docs/UGENCE_ENTERPRISE_AI_GOVERNANCE_CAPABILITY_PIPELINE.md
Run from the repository root:

    python docs/tooling/capability_pipeline/build.py

Outputs:
    docs/UGENCE_ENTERPRISE_AI_GOVERNANCE_CAPABILITY_PIPELINE_v1.0.docx
    docs/UGENCE_ENTERPRISE_AI_GOVERNANCE_CAPABILITY_PIPELINE_v1.0.pdf   (when LibreOffice is available)

The markdown is converted structurally (headings, paragraphs, lists, tables,
callouts, code blocks, links); the mermaid figures are replaced by vector
diagrams from diagrams.py rendered to PNG with cairosvg.
"""
import os
import re
import subprocess
import sys

import cairosvg
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor, Twips

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.abspath(os.path.join(HERE, "..", "..", ".."))
sys.path.insert(0, HERE)
import diagrams as DG  # noqa: E402

SRC = os.path.join(ROOT, "docs", "UGENCE_ENTERPRISE_AI_GOVERNANCE_CAPABILITY_PIPELINE.md")
VERSION = "1.1"
DATE = "4 September 2026"
BASE = f"UGENCE_ENTERPRISE_AI_GOVERNANCE_CAPABILITY_PIPELINE_v{VERSION}"
OUT_DOCX = os.path.join(ROOT, "docs", BASE + ".docx")
COPYRIGHT = "© 2026 Ugence Labs. All rights reserved."
CLASSIFICATION = "Confidential and proprietary · Shared with prospective partners for evaluation"
TITLE = "Ugence Enterprise AI Governance Capability Pipeline"
SUBTITLE = "Repository-Based Capability Map, Development Status and Competitive Cross-Check"

NAVY = RGBColor(0x1A, 0x17, 0x40)
VIOL = RGBColor(0x2A, 0x21, 0x70)
ACCENT = RGBColor(0x51, 0x45, 0xC7)
GREY = RGBColor(0x55, 0x5B, 0x66)
FAINT = RGBColor(0x9A, 0xA0, 0xB0)
LINK = RGBColor(0x1F, 0x4E, 0xA1)
BODY = RGBColor(0x1F, 0x24, 0x30)

PORTRAIT_W = 210 - 20 - 18  # usable mm
LANDSCAPE_W = 297 - 20 - 18

# ----------------------------------------------------------------------------
# Figures
# ----------------------------------------------------------------------------
FIG_DIR = os.path.join(HERE, "png")
os.makedirs(FIG_DIR, exist_ok=True)
FIG_ORDER = ["sequence", "scenario", "minimum_path", "pipeline"]  # mermaid blocks in source order
FIG_CAPTIONS = {
    "sequence": "Figure 1. The nine-stage governance sequence with its feedback loop.",
    "scenario": "Figure 2. The nine stages applied to the governed cloud-scaling scenario.",
    "minimum_path": "Figure 3. Minimum production path for one governed scaling action.",
    "pipeline": "Figure 4. Canonical development pipeline with the stage tags placed on it.",
}
FIG_WIDTH_MM = {"sequence": 165, "scenario": 165, "minimum_path": 120, "pipeline": 168}
FIGS = {}
for key in FIG_ORDER:
    svg, w, h = DG.ALL[key]()
    png = os.path.join(FIG_DIR, key + ".png")
    cairosvg.svg2png(bytestring=svg.encode(), write_to=png, scale=2.6, background_color="white")
    FIGS[key] = png

# ----------------------------------------------------------------------------
# Markdown parsing
# ----------------------------------------------------------------------------
INLINE_RE = re.compile(
    r"(`[^`]+`)"  # code
    r"|(\*\*.+?\*\*)"  # bold
    r"|(\[[^\]]+\]\((https?://[^)\s]+)\))"  # link
    r"|((?<![\w*])\*(?!\*)[^*\n]+?\*(?![\w*]))"  # italic
)


def parse_inline(text, base=None):
    """Return a list of (text, attrs) runs; attrs keys: b, i, code, url, br."""
    base = dict(base or {})
    runs = []
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            runs.append((text[pos : m.start()], dict(base)))
        tok = m.group(0)
        if m.group(1):
            runs.append((tok[1:-1], {**base, "code": True}))
        elif m.group(2):
            runs.extend(parse_inline(tok[2:-2], {**base, "b": True}))
        elif m.group(3):
            label = tok[1 : tok.index("](")]
            runs.extend(parse_inline(label, {**base, "url": m.group(4)}))
        elif m.group(5):
            runs.extend(parse_inline(tok[1:-1], {**base, "i": True}))
        pos = m.end()
    if pos < len(text):
        runs.append((text[pos:], dict(base)))
    return runs


def parse_markdown(md):
    lines = md.split("\n")
    blocks = []
    i = 0
    n = len(lines)
    para = []

    def flush_para():
        if para:
            # lines ending with two spaces are hard line breaks
            segs = []
            for ln in para:
                segs.append((ln.rstrip(), ln.endswith("  ")))
            blocks.append({"t": "p", "segs": segs})
            para.clear()

    while i < n:
        ln = lines[i]
        s = ln.strip()
        if s.startswith("```"):
            flush_para()
            lang = s[3:].strip()
            body = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                body.append(lines[i])
                i += 1
            blocks.append({"t": "code", "lang": lang, "body": body})
            i += 1
            continue
        if not s:
            flush_para()
            i += 1
            continue
        m = re.match(r"^(#{1,4})\s+(.*)$", s)
        if m:
            flush_para()
            blocks.append({"t": "h", "level": len(m.group(1)), "text": m.group(2).strip()})
            i += 1
            continue
        if s == "---":
            flush_para()
            blocks.append({"t": "hr"})
            i += 1
            continue
        if s.startswith("> "):
            flush_para()
            q = []
            while i < n and lines[i].strip().startswith(">"):
                q.append(lines[i].strip()[1:].strip())
                i += 1
            blocks.append({"t": "quote", "text": " ".join(q)})
            continue
        if s.startswith("|"):
            flush_para()
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append(lines[i].strip())
                i += 1
            cells = [[c.strip() for c in r.strip("|").split("|")] for r in rows]
            align = []
            if len(cells) > 1 and all(re.match(r"^:?-{3,}:?$", c) for c in cells[1]):
                align = ["right" if c.endswith(":") and not c.startswith(":") else "left" for c in cells[1]]
                header, body = cells[0], cells[2:]
            else:
                header, body = cells[0], cells[1:]
            blocks.append({"t": "table", "header": header, "rows": body, "align": align})
            continue
        m = re.match(r"^(\s*)([-*])\s+(.*)$", ln)
        if m:
            flush_para()
            items = []
            while i < n:
                mm = re.match(r"^(\s*)([-*])\s+(.*)$", lines[i])
                if not mm:
                    break
                items.append((len(mm.group(1)) // 2, mm.group(3).strip()))
                i += 1
            blocks.append({"t": "ul", "items": items})
            continue
        m = re.match(r"^(\s*)(\d+)\.\s+(.*)$", ln)
        if m:
            flush_para()
            items = []
            while i < n:
                mm = re.match(r"^(\s*)(\d+)\.\s+(.*)$", lines[i])
                if not mm:
                    break
                items.append((len(mm.group(1)) // 2, mm.group(3).strip()))
                i += 1
            blocks.append({"t": "ol", "items": items})
            continue
        para.append(ln)
        i += 1
    flush_para()
    return blocks


# ----------------------------------------------------------------------------
# DOCX helpers
# ----------------------------------------------------------------------------
doc = Document()
styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(10.5)
normal.font.color.rgb = BODY
normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

for lvl, size, color, before, after in [
    (1, 19, NAVY, 18, 8),
    (2, 14, VIOL, 14, 5),
    (3, 11.5, VIOL, 10, 4),
    (4, 10.5, VIOL, 8, 3),
]:
    hs = styles[f"Heading {lvl}"]
    hs.font.name = "Georgia" if lvl <= 2 else "Arial"
    hs.font.size = Pt(size)
    hs.font.bold = True
    hs.font.color.rgb = color
    hs.element.rPr.rFonts.set(qn("w:eastAsia"), hs.font.name)
    hs.paragraph_format.space_before = Pt(before)
    hs.paragraph_format.space_after = Pt(after)
    hs.paragraph_format.keep_with_next = True

for name in ("List Bullet", "List Number"):
    st = styles[name]
    st.font.name = "Arial"
    st.font.size = Pt(10.5)
    st.paragraph_format.space_after = Pt(3)


def add_field(run, instr):
    f1 = OxmlElement("w:fldChar")
    f1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = instr
    f2 = OxmlElement("w:fldChar")
    f2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.text = "1"
    f3 = OxmlElement("w:fldChar")
    f3.set(qn("w:fldCharType"), "end")
    for el in (f1, it, f2, t, f3):
        run._r.append(el)


def shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)


def cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for k, v in (("top", top), ("start", left), ("bottom", bottom), ("end", right)):
        el = OxmlElement(f"w:{k}")
        el.set(qn("w:w"), str(v))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def set_borders(table, color="D8DEE9", sz=4):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    anchor = None
    for tag in ("w:shd", "w:tblLayout", "w:tblCellMar", "w:tblLook"):
        anchor = tblPr.find(qn(tag))
        if anchor is not None:
            break
    if anchor is not None:
        anchor.addprevious(borders)
    else:
        tblPr.append(borders)


def no_borders(table):
    set_borders(table, color="FFFFFF", sz=0)
    tblPr = table._tbl.tblPr
    for b in tblPr.findall(qn("w:tblBorders")):
        for el in b:
            el.set(qn("w:val"), "nil")


def repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    trPr.append(el)


def no_split(row):
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:cantSplit")
    el.set(qn("w:val"), "true")
    trPr.append(el)


def add_hyperlink(paragraph, url, text, size=None, bold=False, italic=False):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rf = OxmlElement("w:rFonts")
    rf.set(qn("w:ascii"), "Arial")
    rf.set(qn("w:hAnsi"), "Arial")
    rPr.append(rf)
    if bold:
        rPr.append(OxmlElement("w:b"))
    if italic:
        rPr.append(OxmlElement("w:i"))
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "1F4EA1")
    rPr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    if size:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(size * 2)))
        rPr.append(sz)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    new_run.append(t)
    hl.append(new_run)
    paragraph._p.append(hl)


def add_runs(p, runs, size=None, color=None):
    for text, a in runs:
        if not text:
            continue
        if a.get("url"):
            add_hyperlink(p, a["url"], text, size=size, bold=a.get("b", False), italic=a.get("i", False))
            continue
        r = p.add_run(text)
        if a.get("b"):
            r.bold = True
        if a.get("i"):
            r.italic = True
        if a.get("code"):
            r.font.name = "Consolas"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            r.font.size = Pt((size or 10.5) - 1)
            r.font.color.rgb = RGBColor(0x2A, 0x21, 0x70)
        elif size:
            r.font.size = Pt(size)
        if color is not None and not a.get("code"):
            r.font.color.rgb = color


def new_number_instance():
    """Create a fresh w:num for the List Number abstract definition, restarting at 1."""
    numbering = doc.part.numbering_part.numbering_definitions._numbering
    style = doc.styles["List Number"]
    num_id = style.element.pPr.numPr.numId.val
    abstract_id = numbering.num_having_numId(num_id).abstractNumId.val
    num = numbering.add_num(abstract_id)
    num.add_lvlOverride(ilvl=0).add_startOverride(1)
    return num.numId


def set_num(p, num_id):
    numPr = p._p.get_or_add_pPr().get_or_add_numPr()
    numPr.get_or_add_ilvl().val = 0
    numPr.get_or_add_numId().val = num_id


def add_inline_paragraph(text, style=None, size=None, color=None, after=None, align=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    add_runs(p, parse_inline(text), size=size, color=color)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    return p


def callout(text, fill="F1F0FB", bar="5145C7", italic=True, size=11, color=VIOL, width_mm=None):
    width_mm = width_mm or CUR_WIDTH[0]
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    c = tbl.cell(0, 0)
    c.width = Mm(width_mm)
    # left accent bar (tcBorders must precede shd and tcMar in tcPr)
    tcPr = c._tc.get_or_add_tcPr()
    b = OxmlElement("w:tcBorders")
    l = OxmlElement("w:left")
    l.set(qn("w:val"), "single")
    l.set(qn("w:sz"), "24")
    l.set(qn("w:space"), "0")
    l.set(qn("w:color"), bar)
    b.append(l)
    tcPr.append(b)
    shade(c, fill)
    cell_margins(c, 120, 120, 180, 160)
    no_borders(tbl)
    p = c.paragraphs[0]
    runs = parse_inline(text, {"i": italic})
    add_runs(p, runs, size=size, color=color)
    p.paragraph_format.space_after = Pt(0)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def code_block(lines, width_mm=None):
    width_mm = width_mm or CUR_WIDTH[0]
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    c = tbl.cell(0, 0)
    c.width = Mm(width_mm)
    shade(c, "F5F7FB")
    cell_margins(c, 100, 100, 160, 160)
    set_borders(tbl, "D8DEE9", 4)
    first = True
    for ln in lines:
        p = c.paragraphs[0] if first else c.add_paragraph()
        first = False
        r = p.add_run(ln if ln.strip() else " ")
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        r.font.size = Pt(9)
        r.font.color.rgb = BODY
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def figure(key):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(FIGS[key], width=Mm(min(FIG_WIDTH_MM[key], CUR_WIDTH[0])))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(FIG_CAPTIONS[key])
    r.italic = True
    r.font.size = Pt(8.8)
    r.font.color.rgb = GREY
    cap.paragraph_format.space_after = Pt(10)


def col_widths(header, rows, total_mm):
    ncol = len(header)
    lens = []
    for ci in range(ncol):
        vals = [len(re.sub(r"[`*\[\]()]", "", header[ci]))]
        for r in rows:
            if ci < len(r):
                txt = re.sub(r"\]\([^)]*\)", "]", r[ci])
                vals.append(len(re.sub(r"[`*]", "", txt)))
        # weight long cells less so a single huge cell does not swallow the table
        vals_sorted = sorted(vals)
        p80 = vals_sorted[int(len(vals_sorted) * 0.8) - 1] if len(vals_sorted) > 1 else vals_sorted[0]
        lens.append(max(3, min(p80, 160)))
    tot = sum(lens)
    widths = [max(7.0, total_mm * l / tot) for l in lens]
    scale = total_mm / sum(widths)
    return [w * scale for w in widths]


def table(block):
    header, rows, align = block["header"], block["rows"], block["align"]
    ncol = len(header)
    widths = col_widths(header, rows, CUR_WIDTH[0])
    fs = 9.0 if ncol <= 3 else 8.4
    tbl = doc.add_table(rows=1, cols=ncol)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    set_borders(tbl)
    hdr = tbl.rows[0]
    repeat_header(hdr)
    for ci, h in enumerate(header):
        c = hdr.cells[ci]
        c.width = Mm(widths[ci])
        shade(c, "1E2340")
        cell_margins(c)
        p = c.paragraphs[0]
        add_runs(p, parse_inline(h, {"b": True}), size=fs, color=RGBColor(0xFF, 0xFF, 0xFF))
        p.paragraph_format.space_after = Pt(0)
        if align and align[ci] == "right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for ri, row in enumerate(rows):
        tr = tbl.add_row()
        no_split(tr)
        for ci in range(ncol):
            c = tr.cells[ci]
            c.width = Mm(widths[ci])
            if ri % 2 == 1:
                shade(c, "F5F7FB")
            cell_margins(c)
            txt = row[ci] if ci < len(row) else ""
            p = c.paragraphs[0]
            add_runs(p, parse_inline(txt), size=fs)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            if align and align[ci] == "right":
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ----------------------------------------------------------------------------
# Sections, header and footer
# ----------------------------------------------------------------------------
CUR_WIDTH = [PORTRAIT_W]


def setup_section(sec, landscape=False, first_page_blank=False):
    if landscape:
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_width, sec.page_height = Mm(297), Mm(210)
    else:
        sec.orientation = WD_ORIENT.PORTRAIT
        sec.page_width, sec.page_height = Mm(210), Mm(297)
    sec.left_margin, sec.right_margin = Mm(20), Mm(18)
    sec.top_margin, sec.bottom_margin = Mm(18), Mm(16)
    sec.header_distance, sec.footer_distance = Mm(9), Mm(8)
    width = LANDSCAPE_W if landscape else PORTRAIT_W
    CUR_WIDTH[0] = width
    sec.different_first_page_header_footer = first_page_blank
    sec.header.is_linked_to_previous = False
    sec.footer.is_linked_to_previous = False
    # header: title left, classification right
    hp = sec.header.paragraphs[0]
    hp.text = ""
    for tw in (4680, 9360):
        hp.paragraph_format.tab_stops.add_tab_stop(Twips(tw), WD_TAB_ALIGNMENT.CLEAR)
    hp.paragraph_format.tab_stops.add_tab_stop(Mm(width), WD_TAB_ALIGNMENT.RIGHT)
    r = hp.add_run(f"{TITLE}")
    r.font.size = Pt(8)
    r.font.color.rgb = FAINT
    r = hp.add_run(f"\tVersion {VERSION} · {DATE}")
    r.font.size = Pt(8)
    r.font.color.rgb = FAINT
    para_border(hp, "bottom")
    # footer: copyright left, page x of y right
    fp = sec.footer.paragraphs[0]
    fp.text = ""
    for tw in (4680, 9360):
        fp.paragraph_format.tab_stops.add_tab_stop(Twips(tw), WD_TAB_ALIGNMENT.CLEAR)
    fp.paragraph_format.tab_stops.add_tab_stop(Mm(width), WD_TAB_ALIGNMENT.RIGHT)
    r = fp.add_run(f"{COPYRIGHT}  ·  Confidential and proprietary  ·  For partnership evaluation only")
    r.font.size = Pt(8)
    r.font.color.rgb = FAINT
    r = fp.add_run("\tPage ")
    r.font.size = Pt(8)
    r.font.color.rgb = FAINT
    r = fp.add_run()
    r.font.size = Pt(8)
    r.font.color.rgb = FAINT
    add_field(r, "PAGE")
    r = fp.add_run(" of ")
    r.font.size = Pt(8)
    r.font.color.rgb = FAINT
    r = fp.add_run()
    r.font.size = Pt(8)
    r.font.color.rgb = FAINT
    add_field(r, "NUMPAGES")
    para_border(fp, "top")


def para_border(p, edge, color="D8DEE9"):
    pPr = p._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    el = OxmlElement(f"w:{edge}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), "4")
    el.set(qn("w:space"), "1")
    el.set(qn("w:color"), color)
    bdr.append(el)
    tabs = pPr.find(qn("w:tabs"))
    if tabs is not None:
        tabs.addprevious(bdr)
    else:
        pPr.append(bdr)


def new_section(landscape):
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    setup_section(sec, landscape=landscape)
    return sec


# ----------------------------------------------------------------------------
# Cover sheet and document control
# ----------------------------------------------------------------------------
sec0 = doc.sections[0]
setup_section(sec0, landscape=False, first_page_blank=True)
# first-page footer carries only the copyright line (no page number)
ffp = sec0.first_page_footer.paragraphs[0]
ffp.text = ""
r = ffp.add_run(COPYRIGHT)
r.font.size = Pt(8)
r.font.color.rgb = FAINT
sec0.first_page_header.paragraphs[0].text = ""


def cover_par(text, size, color, bold=False, italic=False, after=4, font=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    if font:
        run.font.name = font
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    return p


# violet band at the top of the cover
band = doc.add_table(rows=1, cols=1)
band.alignment = WD_TABLE_ALIGNMENT.CENTER
band.autofit = False
bc = band.cell(0, 0)
bc.width = Mm(PORTRAIT_W)
shade(bc, "16143A")
cell_margins(bc, 360, 360, 320, 320)
no_borders(band)
bp = bc.paragraphs[0]
r = bp.add_run("UGENCE LABS  ·  ENTERPRISE AI GOVERNANCE PLATFORM")
r.font.size = Pt(9.5)
r.bold = True
r.font.color.rgb = RGBColor(0x9A, 0xA0, 0xD8)
bp.paragraph_format.space_after = Pt(14)
bp2 = bc.add_paragraph()
r = bp2.add_run(TITLE)
r.font.size = Pt(27)
r.bold = True
r.font.name = "Georgia"
r._element.rPr.rFonts.set(qn("w:eastAsia"), "Georgia")
r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
bp2.paragraph_format.space_after = Pt(8)
bp3 = bc.add_paragraph()
r = bp3.add_run(SUBTITLE)
r.font.size = Pt(14)
r.bold = True
r.font.color.rgb = RGBColor(0xC9, 0xCF, 0xEA)
bp3.paragraph_format.space_after = Pt(4)
bp4 = bc.add_paragraph()
r = bp4.add_run("Define → Propose → Verify → Decide → Authorize → Clear → Execute → Assure → Measure")
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x9A, 0xA0, 0xD8)
bp4.paragraph_format.space_after = Pt(0)

doc.add_paragraph().paragraph_format.space_after = Pt(14)
cover_par(
    "Agents and models may propose. Independent governance components verify, decide, authorize and clear. "
    "The runtime executes only within that authority, while assurance observes the trajectory and outcome.",
    12,
    VIOL,
    italic=True,
    after=18,
    font="Georgia",
)
cover_par("Prepared by Ugence Labs", 13, NAVY, bold=True, after=1)
cover_par(f"ugence.ai  ·  Version {VERSION}  ·  {DATE}", 11, GREY, after=16)

# document control table
ctrl = [
    ("Document", TITLE),
    ("Edition", f"Version {VERSION} · Partner evaluation edition"),
    ("Date", DATE),
    ("Source of truth", "docs/UGENCE_ENTERPRISE_AI_GOVERNANCE_CAPABILITY_PIPELINE.md in rasaha/symbolu (merged in PR #1584, commit 8c6e5ec6)"),
    ("Scope", "45 platform capabilities under packages/; the two packaged business-solution examples are excluded"),
    ("Classification", CLASSIFICATION),
    ("Intended recipients", "Prospective clients and development partners evaluating a partnership with Ugence Labs"),
    ("Owner", "Ugence Labs"),
]
t = doc.add_table(rows=0, cols=2)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.autofit = False
set_borders(t, "D8DEE9", 4)
for k, v in ctrl:
    row = t.add_row()
    c0, c1 = row.cells
    c0.width, c1.width = Mm(38), Mm(PORTRAIT_W - 38)
    shade(c0, "F5F7FB")
    cell_margins(c0)
    cell_margins(c1)
    rr = c0.paragraphs[0].add_run(k)
    rr.bold = True
    rr.font.size = Pt(9)
    rr.font.color.rgb = VIOL
    rr = c1.paragraphs[0].add_run(v)
    rr.font.size = Pt(9)
    for cc in (c0, c1):
        cc.paragraphs[0].paragraph_format.space_after = Pt(0)
doc.add_paragraph().paragraph_format.space_after = Pt(10)
cover_par(
    "Repository-based capability map grounded in the inspected codebase snapshot. Contract-only and research-only packages "
    "are described according to their intended position while their present limitations are stated explicitly. "
    "Nothing in this document represents a production deployment, a customer validation or legal advice.",
    9.5,
    GREY,
    italic=True,
    after=8,
)
cover_par(
    f"{COPYRIGHT} This document and the architecture it describes are the confidential and proprietary property of "
    "Ugence Labs. It is provided to the recipient solely to evaluate a development partnership with Ugence Labs and may be "
    "shared within the recipient's evaluation team for that purpose only. It may not otherwise be reproduced, distributed "
    "or disclosed without prior written permission. It is not an offer and creates no obligation on either party. Product "
    "and company names cited as competitor analogues are trademarks of their respective owners and are referenced for "
    "comparison only.",
    8.5,
    FAINT,
    after=0,
)

# ---- revision history + contents page ----
doc.add_page_break()
doc.add_heading("Document control", level=1)
rev = doc.add_table(rows=1, cols=4)
rev.alignment = WD_TABLE_ALIGNMENT.CENTER
rev.autofit = False
set_borders(rev)
rw = [22, 30, 40, PORTRAIT_W - 92]
for ci, h in enumerate(["Version", "Date", "Author", "Change"]):
    c = rev.rows[0].cells[ci]
    c.width = Mm(rw[ci])
    shade(c, "1E2340")
    cell_margins(c)
    rr = c.paragraphs[0].add_run(h)
    rr.bold = True
    rr.font.size = Pt(9)
    rr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    c.paragraphs[0].paragraph_format.space_after = Pt(0)
for vals in [
    ("0.1", "2026-09-04", "Ugence Labs", "Capability map (Sections 1–15) and Appendix A use case, as supplied."),
    ("0.2", "2026-09-04", "Ugence Labs", "Appendix B added: development status of the 45 capabilities with stage tags, evidence and canonical pipeline diagram."),
    ("0.3", "2026-09-04", "Ugence Labs", "Body text cross-checked against package source; five claims corrected in place; competitor supplement and Appendix C added."),
    ("1.0", "2026-09-04", "Ugence Labs", "First docx edition with cover sheet, document control, figures and copyright footer."),
    ("1.1", "2026-09-04", "Ugence Labs", "Partner evaluation edition: confidentiality wording aligned to distribution to prospective clients and partners."),
]:
    row = rev.add_row()
    for ci, v in enumerate(vals):
        c = row.cells[ci]
        c.width = Mm(rw[ci])
        cell_margins(c)
        rr = c.paragraphs[0].add_run(v)
        rr.font.size = Pt(9)
        c.paragraphs[0].paragraph_format.space_after = Pt(0)
doc.add_page_break()
doc.add_heading("Contents", level=1)
# Static table of contents. Page numbers come from a first-pass PDF render
# (docs/tooling/capability_pipeline/toc_pages.json); the entry count is fixed so
# a second pass reproduces the same pagination.
import json  # noqa: E402

TOC_PAGES_FILE = os.path.join(HERE, "toc_pages.json")
TOC_PAGES = json.load(open(TOC_PAGES_FILE)) if os.path.exists(TOC_PAGES_FILE) else {}
md_for_toc = open(SRC, encoding="utf-8").read()
TOC_ENTRIES = []
for ln in md_for_toc.split("\n"):
    m = re.match(r"^(#{1,2})\s+(.*)$", ln)
    if m and not ln.startswith("# Ugence Enterprise"):
        TOC_ENTRIES.append((len(m.group(1)), m.group(2).strip()))
for lvl, text in TOC_ENTRIES:
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(Mm(PORTRAIT_W), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    p.paragraph_format.space_after = Pt(2 if lvl == 2 else 4)
    p.paragraph_format.space_before = Pt(6 if lvl == 1 else 0)
    p.paragraph_format.left_indent = Mm(0 if lvl == 1 else 6)
    r = p.add_run(text)
    r.font.size = Pt(10.5 if lvl == 1 else 9.5)
    r.bold = lvl == 1
    r.font.color.rgb = NAVY if lvl == 1 else BODY
    r2 = p.add_run("\t" + str(TOC_PAGES.get(text, "")))
    r2.font.size = Pt(10.5 if lvl == 1 else 9.5)
    r2.font.color.rgb = GREY
doc.add_page_break()
settings = doc.settings.element
uf = OxmlElement("w:updateFields")
uf.set(qn("w:val"), "true")
_compat = settings.find(qn("w:compat"))
if _compat is not None:
    _compat.addprevious(uf)
else:
    settings.append(uf)
_zoom = settings.find(qn("w:zoom"))
if _zoom is not None and _zoom.get(qn("w:percent")) is None:
    _zoom.set(qn("w:percent"), "100")

# ----------------------------------------------------------------------------
# Body
# ----------------------------------------------------------------------------
md = open(SRC, encoding="utf-8").read()
blocks = parse_markdown(md)

# drop the markdown title and the two bold metadata lines (they live on the cover)
assert blocks[0]["t"] == "h" and blocks[0]["level"] == 1
blocks = blocks[1:]
if blocks[0]["t"] == "p" and blocks[0]["segs"][0][0].startswith("**Repository-based"):
    blocks = blocks[1:]

fig_iter = iter(FIG_ORDER)
landscape = False
first_h1 = True


def wide_table_ahead(idx, horizon=3):
    """True when a >=5-column table follows within `horizon` blocks with only headings/paragraphs between."""
    for j in range(idx, min(idx + horizon + 1, len(blocks))):
        bj = blocks[j]
        if bj["t"] == "table":
            return len(bj["header"]) >= 5
        if bj["t"] not in ("h", "p"):
            return False
        if bj["t"] == "h" and bj["level"] == 1:
            return False
    return False


i = 0
while i < len(blocks):
    b = blocks[i]
    t = b["t"]
    if t in ("h", "p") and not landscape and wide_table_ahead(i):
        new_section(True)
        landscape = True
    if t == "h":
        lvl = b["level"]
        if lvl == 1:
            # part-level heading: fresh portrait page
            if landscape:
                new_section(False)
                landscape = False
            elif not first_h1:
                doc.add_page_break()
            first_h1 = False
            doc.add_heading(b["text"], level=1)
        else:
            doc.add_heading(b["text"], level=min(lvl, 4))
    elif t == "p":
        p = doc.add_paragraph()
        for si, (seg, br) in enumerate(b["segs"]):
            add_runs(p, parse_inline(seg))
            if br and si < len(b["segs"]) - 1:
                p.add_run().add_break(WD_BREAK.LINE)
            elif si < len(b["segs"]) - 1:
                p.add_run(" ")
    elif t == "quote":
        callout(b["text"])
    elif t == "hr":
        # the closing repository note follows; render it as a note callout
        nxt = blocks[i + 1] if i + 1 < len(blocks) else None
        if nxt and nxt["t"] == "p":
            callout(" ".join(s for s, _ in nxt["segs"]), fill="EEF6F1", bar="2E7D5B", italic=False, size=9.8, color=BODY)
            i += 1
    elif t == "code":
        if b["lang"] == "mermaid":
            figure(next(fig_iter))
        else:
            code_block(b["body"])
    elif t == "ul":
        for lvl_, txt in b["items"]:
            style = "List Bullet" if lvl_ == 0 else "List Bullet 2"
            add_inline_paragraph(txt, style=style)
    elif t == "ol":
        nid = new_number_instance()
        for lvl_, txt in b["items"]:
            p = add_inline_paragraph(txt, style="List Number")
            set_num(p, nid)
    elif t == "table":
        wide = len(b["header"]) >= 5
        if wide and not landscape:
            new_section(True)
            landscape = True
        table(b)
        if landscape:
            # return to portrait immediately after a wide table
            new_section(False)
            landscape = False
    i += 1

# restart numbering is not needed (List Number continues); acceptable for this document.

# core properties
cp = doc.core_properties
cp.title = TITLE
cp.subject = SUBTITLE
cp.author = "Ugence Labs"
cp.keywords = "Ugence, AI governance, capability pipeline, ActionGate, Action Clearance, Risk Authority, Decision Authority, Agent Runtime, development status, competitor analysis"
cp.comments = COPYRIGHT
cp.category = "Architecture"
cp.version = VERSION

doc.save(OUT_DOCX)
# stamp company into app.xml
import zipfile, shutil  # noqa: E402

tmp = OUT_DOCX + ".tmp"
with zipfile.ZipFile(OUT_DOCX) as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
    for item in zin.infolist():
        data = zin.read(item.filename)
        if item.filename == "docProps/app.xml":
            s = data.decode("utf-8")
            if "<Company/>" in s:
                s = s.replace("<Company/>", "<Company>Ugence Labs</Company>", 1)
            elif "<Company>" in s:
                s = re.sub(r"<Company>.*?</Company>", "<Company>Ugence Labs</Company>", s, count=1)
            else:
                s = s.replace("<LinksUpToDate>", "<Company>Ugence Labs</Company><LinksUpToDate>", 1)
            data = s.encode("utf-8")
        zout.writestr(item, data)
shutil.move(tmp, OUT_DOCX)
print("DOCX written:", OUT_DOCX)

# PDF via LibreOffice when available
soffice = shutil.which("soffice")
if soffice and "--no-pdf" not in sys.argv:
    import tempfile

    profile = tempfile.mkdtemp(prefix="lo_profile_")
    subprocess.run(
        [soffice, f"-env:UserInstallation=file://{profile}", "--headless", "--convert-to", "pdf", "--outdir", os.path.dirname(OUT_DOCX), OUT_DOCX],
        check=False,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        timeout=600,
    )
    pdf = OUT_DOCX[:-5] + ".pdf"
    print("PDF written:" if os.path.exists(pdf) else "PDF not produced:", pdf)
